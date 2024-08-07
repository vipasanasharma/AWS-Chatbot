import os
import boto3
import uuid
from docx import Document as DocxDocument
from langchain.embeddings import BedrockEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.schema import Document

# Load environment variables from .env file
AWS_REGION = os.getenv('AWS_REGIONS', 'us-east-1')
BEDROCK_MODEL_ID = os.getenv('BEDROCK_MODEL_ID', 'amazon.titan-embed-text-v1')
OUTPUT_BUCKET_NAME = os.getenv('S3_BUCKET_NAME', 'mirror-embeddings')
INPUT_BUCKET_NAME = 'compliance-month'

# Initialize AWS clients and variables
s3_client = boto3.client("s3", region_name=AWS_REGION)
bedrock_client = boto3.client(service_name="bedrock-runtime", region_name=AWS_REGION)
bedrock_embeddings = BedrockEmbeddings(model_id="amazon.titan-embed-text-v1", client=bedrock_client)

def get_unique_id():
    return str(uuid.uuid4())

def split_text(text, chunk_size, chunk_overlap):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    texts = text_splitter.split_text(text)
    documents = [Document(page_content=text) for text in texts]
    return documents

def create_vector_store(request_id, documents):
    vectorstore_faiss = FAISS.from_documents(documents, bedrock_embeddings)
    file_name = f"{request_id}.bin"
    folder_path = "/tmp/"
    vectorstore_faiss.save_local(index_name=file_name, folder_path=folder_path)

    # Upload files to S3
    s3_client.upload_file(Filename=folder_path + file_name + ".faiss", Bucket=OUTPUT_BUCKET_NAME, Key=f"{request_id}.faiss")
    s3_client.upload_file(Filename=folder_path + file_name + ".pkl", Bucket=OUTPUT_BUCKET_NAME, Key=f"{request_id}.pkl")

    return True

def authenticate_user():
    # Simulate basic authentication
    authenticated = True  # Replace with your actual authentication logic
    
    if authenticated:
        print("User authenticated successfully.")
    else:
        print("Authentication failed. Access denied.")
        raise PermissionError("Authentication failed.")

def process_docx(docx_path):
    authenticate_user()  # Check authentication before proceeding

    request_id = get_unique_id()
    document = DocxDocument(docx_path)

    # Extract text from the document
    full_text = []
    for para in document.paragraphs:
        full_text.append(para.text)
    text = '\n'.join(full_text)

    # Split text into chunks
    splitted_docs = split_text(text, 1000, 200)

    # Create vector store
    result = create_vector_store(request_id, splitted_docs)

    return result

def download_docx_from_s3(bucket_name, key, download_path):
    s3_client.download_file(Bucket=bucket_name, Key=key, Filename=download_path)

def main():
    print("This is the DOCX processing script")
    
    # List DOCX files in the input S3 bucket
    response = s3_client.list_objects_v2(Bucket=INPUT_BUCKET_NAME)
    
    for obj in response.get('Contents', []):
        if obj['Key'].endswith('.docx'):
            docx_key = obj['Key']
            download_path = f"/tmp/{docx_key.split('/')[-1]}"
            
            # Download DOCX file
            download_docx_from_s3(INPUT_BUCKET_NAME, docx_key, download_path)
            
            # Process DOCX file
            process_docx(download_path)
            print(f"Processed and uploaded embeddings for: {docx_key}")

if __name__ == "__main__":
    main()
